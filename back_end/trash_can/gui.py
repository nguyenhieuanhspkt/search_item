import streamlit as st
import pandas as pd
import os
import shutil
import io
from docx import Document  # Thư viện xử lý file Word
from engine import HybridSearchEngine
from whoosh.index import create_in
from whoosh.fields import Schema, TEXT, ID

# 1. CẤU HÌNH TRANG
st.set_page_config(page_title="Hệ thống Thẩm định Vật tư", page_icon="🛡️", layout="wide")

# 2. HÀM NẠP DỮ LIỆU (INDEXING) - ĐÃ CẬP NHẬT TRƯỜNG DVT
def build_index(df, index_dir="vattu_index"):
    if os.path.exists(index_dir):
        shutil.rmtree(index_dir)
    os.makedirs(index_dir)

    # Schema mới bao gồm trường dvt (Đơn vị tính)
    schema = Schema(
        ma_vattu=ID(stored=True),
        ma_erp=ID(stored=True),
        ten_vattu=TEXT(stored=True),
        thong_so=TEXT(stored=True),
        dvt=TEXT(stored=True),      # THÊM DÒNG NÀY
        all_text=TEXT(stored=True)
    )

    ix = create_in(index_dir, schema)
    writer = ix.writer()
    df = df.fillna("")

    for _, row in df.iterrows():
        # Ánh xạ thông minh các tên cột có thể có trong file Excel
        ma = str(row.get('Mã vật tư', row.get('ma_vattu', '')))
        ten = str(row.get('Tên vật tư', row.get('ten_vattu', '')))
        ts = str(row.get('Mã hiệu/Thông số kỹ thuật', row.get('thong_so', '')))
        erp = str(row.get('Mã ERP', row.get('Ghi chú', row.get('ma_erp', ''))))
        
        # Ánh xạ Đơn vị tính từ Excel
        dvt_val = str(row.get('Đơn vị tính', row.get('ĐVT', row.get('dvt', ''))))
        
        # Xử lý làm sạch dữ liệu
        ts = "" if ts.lower() == "none" else ts
        
        # Gom text để AI tìm kiếm (bao gồm cả ĐVT để tăng độ chính xác)
        full_info = f"{ten} {ts} {dvt_val} {ma} {erp}"
        
        writer.add_document(
            ma_vattu=ma,
            ma_erp=erp,
            ten_vattu=ten,
            thong_so=ts,
            dvt=dvt_val,            # THÊM DÒNG NÀY
            all_text=full_info
        )
    writer.commit()

# 3. KHỞI TẠO ENGINE
@st.cache_resource
def load_engine():
    model_path = 'keepitreal/vietnamese-sbert'
    index_dir = "vattu_index"
    try:
        if os.path.exists(index_dir) and len(os.listdir(index_dir)) > 0:
            return HybridSearchEngine(model_path, index_dir), "Sẵn sàng"
        return None, "Chưa có dữ liệu"
    except Exception as e:
        return None, f"Lỗi: {str(e)}"

# --- GIAO DIỆN CHÍNH ---
st.title("🛡️ Hệ thống Thẩm định Vật tư - Tổ Thẩm định")

tab_search, tab_batch, tab_admin = st.tabs([
    "🔍 Tìm kiếm đơn lẻ", 
    "📑 Thẩm định File Word", 
    "⚙️ Quản trị dữ liệu"
])

# --- TAB 1: TÌM KIẾM ĐƠN LẺ ---
with tab_search:
    engine, status = load_engine()
    if status != "Sẵn sàng":
        st.warning("Hệ thống chưa có dữ liệu. Hãy sang tab 'Quản trị dữ liệu' để nạp file Excel.")
    else:
        query = st.text_area("Nhập tên hoặc thông số cần thẩm định:", height=100)
        if st.button("🚀 THẨM ĐỊNH") and query:
            with st.spinner("AI đang phân tích..."):
                results = engine.search(query)
                if not results:
                    st.error("Không tìm thấy kết quả.")
                else:
                    for i, res in enumerate(results):
                        # Hiển thị kết quả bao gồm Đơn vị tính (dvt)
                        label = f"Top {i+1}: {res['ten']} | ĐVT: {res.get('dvt', 'N/A')} (ERP: {res['erp']})"
                        with st.expander(label):
                            st.write(f"**Mã ERP:** `{res['erp']}` | **Mã nội bộ:** `{res['ma']}`")
                            st.write(f"**Đơn vị tính:** `{res.get('dvt', 'N/A')}`")
                            st.write(f"**Thông số:** {res['ts']}")
                            relevance = min(max((res['ai_relevance'] + 5) / 15, 0.0), 1.0)
                            st.progress(relevance)

# --- TAB 2: THẨM ĐỊNH FILE WORD ---
with tab_batch:
    st.subheader("📑 Thẩm định hàng loạt từ Bảng trong Word")
    st.info("Hệ thống sẽ quét các bảng trong Word và hiển thị tiến trình thực tế.")
    
    word_file = st.file_uploader("Chọn file Word có chứa bảng danh mục", type=["docx"])
    engine, status = load_engine()

    if word_file and status == "Sẵn sàng":
        if st.button("🚀 BẮT ĐẦU QUÉT BẢNG & THẨM ĐỊNH"):
            doc = Document(word_file)
            all_rows_data = []
            
            total_rows = sum(len(table.rows) - 1 for table in doc.tables)
            
            if total_rows == 0:
                st.error("Không tìm thấy bảng hoặc dữ liệu trong file Word.")
            else:
                progress_bar = st.progress(0)
                status_text = st.empty()
                current_count = 0
                
                for table in doc.tables:
                    for i, row in enumerate(table.rows[1:]): # Bỏ qua header
                        current_count += 1
                        cells = row.cells
                        
                        progress_val = current_count / total_rows
                        progress_bar.progress(progress_val)
                        status_text.text(f"⏳ Đang thẩm định dòng {current_count}/{total_rows}...")

                        try:
                            stt = cells[0].text.strip()
                            ten_hang = cells[1].text.strip()
                            thong_so = cells[2].text.strip()
                            dvt_word = cells[3].text.strip()
                        except IndexError:
                            continue 

                        if not ten_hang and not thong_so:
                            continue

                        # Tìm kiếm AI (Search engine tự động lấy trường dvt từ index)
                        query_str = f"{ten_hang} {thong_so}"
                        results = engine.search(query_str, top_k=1)
                        
                        res_info = {
                            "STT": stt,
                            "Tên hàng hóa (Word)": ten_hang,
                            "Thông số (Word)": thong_so,
                            "ĐVT (Word)": dvt_word,
                            "Kết quả tìm thấy (Top 1)": "❌ Không tìm thấy",
                            "Thông số hệ thống": "",
                            "ĐVT hệ thống": "",
                            "Điểm tin cậy": 0
                        }
                        
                        if results:
                            best = results[0]
                            res_info.update({
                                "Kết quả tìm thấy (Top 1)": best['ten'],
                                "Thông số hệ thống": best['ts'],
                                "ĐVT hệ thống": best.get('dvt', 'N/A'), 
                                "Điểm tin cậy": round(best['final_score'], 2)
                            })
                        
                        all_rows_data.append(res_info)

                status_text.success(f"✅ Đã thẩm định xong {total_rows} dòng!")

                if all_rows_data:
                    df_final = pd.DataFrame(all_rows_data)
                    st.write("### Bản xem trước kết quả:")
                    st.dataframe(df_final)

                    output = io.BytesIO()
                    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                        df_final.to_excel(writer, index=False, sheet_name='ThamDinhVattu')
                    
                    st.download_button(
                        label="📥 Tải Báo cáo Thẩm định (.xlsx)",
                        data=output.getvalue(),
                        file_name="Bao_cao_Tham_dinh_Vat_tu.xlsx",
                        mime="application/vnd.ms-excel"
                    )

# --- TAB 3: QUẢN TRỊ DỮ LIỆU ---
with tab_admin:
    st.subheader("Cấu trúc dữ liệu hệ thống")
    uploaded_file = st.file_uploader("Nạp file Excel danh mục gốc", type=["xlsx"])
    if uploaded_file:
        df_preview = pd.read_excel(uploaded_file)
        st.write("Dữ liệu nạp vào (5 dòng đầu):")
        st.dataframe(df_preview.head())
        if st.button("🚀 CẬP NHẬT DATABASE"):
            with st.spinner("Đang thực hiện xây dựng chỉ mục..."):
                build_index(df_preview)
                st.cache_resource.clear()
                st.success("Hệ thống đã cập nhật dữ liệu mới thành công!")

st.sidebar.caption("Tổ Thẩm định - TaskApp v2.6")