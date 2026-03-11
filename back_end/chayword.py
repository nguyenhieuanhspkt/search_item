from docx import Document

# Tạo file Word
doc = Document()
doc.add_heading('Bảng thống kê chi tiết các nội dung thay đổi', level=1)

# Dữ liệu bảng
data = [
["STT","Vị trí điều khoản","Trước điều chỉnh","Sau điều chỉnh"],

["1","Căn cứ (đầu Quyết định)",
"Chưa có căn cứ Quy chế 30",
"Căn cứ Quy chế mua sắm số 30/QĐ-HĐTV ngày 12/02/2026 của Hội đồng thành viên Tập đoàn điện lực Việt Nam"],

["2","Mục lục (Trang)",
"Bước 2: trang 9; Bước 3: trang 9; Bước 4,5: trang 10; Bước 6: trang 17",
"Bước 2: trang 10; Bước 3: trang 11; Bước 4,5: trang 11; Bước 6: trang 18"],

["3","Điều 1, Khoản 2",
"Đảm bảo tính chính xác, công khai, minh bạch, công bằng hiệu quả",
"Đảm bảo tính chính xác, công khai, minh bạch, công bằng hiệu quả và trách nhiệm giải trình"],

["4","Điều 2, Khoản 1",
"Đảm bảo chi phí dự toán được tính đúng, tính đủ",
"Đảm bảo chi phí dự toán được xác định tương ứng phù hợp với nội dung công việc"],

["5","Điều 2, Khoản 2",
"...nguyên tắc công khai, minh bạch, hiệu quả, cạnh tranh, công bằng, minh bạch",
"...nguyên tắc công khai, minh bạch, hiệu quả, cạnh tranh, công bằng, minh bạch và trách nhiệm giải trình"],

["6","Điều 4 (Luật)",
"Luật Đấu thầu số 22/2023/QH15 ngày 23/06/2023",
"Luật Đấu thầu số 22/2023/QH15; Luật số 57/2024/QH15 ngày 29/11/2024; Luật số 90/2025/QH15 ngày 25/6/2025"],

["7","Điều 4 (Nghị định)",
"Nghị định số 24/2024/NĐ-CP ngày 27/02/2024",
"Nghị định số 214/2025/NĐ-CP ngày 04/08/2025 hướng dẫn Luật Đấu thầu về lựa chọn nhà thầu"],

["8","Điều 4 (Văn bản EVN)",
"Quyết định số 127/QĐ-HĐTV ngày 01/10/2021",
"Quy chế mua sắm số 30/QĐ-HĐTV ngày 12/02/2026 của Hội đồng thành viên EVN"],

["9","Điều 4 (Văn bản VT4)",
"Quyết định số 127/QĐ-NĐVT4 ngày 01/10/2023",
"Quyết định số 1884/QĐ-NĐVT4 ngày 27/12/2025"],

["10","Điều 7, Bước 1",
"...về ĐTXD, SXKD, đấu thầu",
"...về ĐTXD, SXKD, mua sắm"],

["11","Điều 7, Bước 1.2",
"Chưa quy định ngoại lệ",
"Bổ sung: Không thẩm định KHLCNT với gói thầu không thuộc dự án đầu tư và mua sắm nhỏ lẻ"],

["12","Điều 7, Bước 1.3",
"Chưa quy định ngoại lệ",
"Bổ sung: Không phải thẩm định với mua sắm nhỏ lẻ; không bắt buộc với chào hàng cạnh tranh, chỉ định thầu, ký hợp đồng trực tiếp"],

["13","Điều 7, Bước 2",
"Chia 2 trường hợp xử lý riêng",
"Gộp chung quy trình tiếp nhận và xử lý đề nghị thẩm định"],

["14","Điều 7, Bước 2",
"Phân công 1 nhân sự thực hiện",
"Phân công 01 hoặc một số nhân sự chủ trì và tối thiểu 01 nhân sự phối hợp"],

["15","Điều 7, Bước 5.1.1",
"Quy chế đầu tư xây dựng; Quy chế SXKD điện",
"Bổ sung: Văn bản phân cấp cho Giám đốc đơn vị trực thuộc (QĐ 630); Quy chế mua sắm (QC 30)"],

["16","Điều 7, Bước 5.2.2",
"Giá gói thầu không bao gồm tùy chọn mua thêm",
"Loại bỏ nội dung tùy chọn mua thêm"],

["17","Điều 7, Bước 5.2.2",
"Nêu rõ nguồn vốn bao gồm SXKD, ĐTXD, Quỹ đầu tư phát triển",
"Bỏ liệt kê chi tiết nguồn vốn"],

["18","Điều 7, Bước 5.2.2",
"Mục Tùy chọn mua thêm (nếu có)",
"Loại bỏ toàn bộ mục này"],

["19","Điều 7, Bước 6.1",
"Quy trình phát hành báo cáo thẩm định chung",
"Bổ sung chi tiết quy trình xử lý ý kiến góp ý và trường hợp 1 nhân sự thực hiện"]
]

# Tạo bảng
table = doc.add_table(rows=1, cols=4)

# Header
for i, col in enumerate(data[0]):
    table.rows[0].cells[i].text = col

# Thêm dữ liệu
for row in data[1:]:
    cells = table.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val

# Ghi chú cuối
doc.add_paragraph(
"Lưu ý: Việc điều chỉnh nhằm cập nhật các văn bản pháp luật mới "
"(Luật Đấu thầu 2023, Nghị định 214/2025/NĐ-CP) và các quy chế nội bộ mới của EVN "
"(Quy chế 30, Quyết định 630) nhằm đảm bảo tính tuân thủ và trách nhiệm giải trình."
)

# Lưu file
doc.save("bang_thong_ke_thay_doi_quy_trinh_moi.docx")

print("Đã tạo file Word thành công.")